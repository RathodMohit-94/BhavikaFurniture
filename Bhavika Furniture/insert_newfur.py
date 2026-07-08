from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

doc = Document(r'C:\Users\miraj\OneDrive\Desktop\folder\MirajPipaliya_Internship_Report_University.docx')

# ── find insertion anchor: last bullet of section 4.9 ────────────────────────
anchor = None
for para in doc.paragraphs:
    if para.text.strip().startswith('N+1 Query Prevention'):
        anchor = para
        break

if anchor is None:
    for para in doc.paragraphs:
        if '4.9' in para.text and 'Performance' in para.text:
            anchor = para
            break

assert anchor is not None, "Could not find insertion anchor paragraph"

# ── content blocks (text, style, is_bold_heading) ────────────────────────────
blocks = [
    # ── 4.10 ─────────────────────────────────────────────────────────────────
    ("4.10 Additional Project: Bhavika Furniture – Full Stack E-Commerce & CMS Platform", "Normal", True),
    (
        "During the internship period, in addition to contributing to the ERP and HRMS system at "
        "Airwix Technologies, I independently developed a complete full-stack web project titled "
        "\"Bhavika Furniture\" — a furniture e-commerce platform with a built-in content management system "
        "(CMS). This project was undertaken as a hands-on personal initiative to apply and "
        "consolidate the full-stack development skills acquired during the internship in a "
        "real-world product scenario. The Bhavika Furniture project is a production-ready web application "
        "consisting of two integrated parts: a public-facing storefront for furniture browsing, "
        "and a dynamic admin panel for store content management.",
        "Normal", False
    ),

    # ── 4.11 ─────────────────────────────────────────────────────────────────
    ("4.11 Problem Identification", "Normal", True),
    (
        "Small furniture retailers and boutique home décor businesses typically face two major "
        "challenges when establishing an online presence. First, building a custom storefront "
        "requires ongoing developer involvement for even simple content changes such as updating "
        "a hero image, adding a new product, or editing a contact address. Business owners become "
        "completely dependent on developers for tasks that should be self-serviceable. Second, "
        "most off-the-shelf e-commerce solutions are either too complex and expensive for small "
        "businesses, or too rigid to accommodate the specific branding and content requirements "
        "of a boutique brand.",
        "Normal", False
    ),
    (
        "The Bhavika Furniture project addresses both problems by combining a polished, brand-forward "
        "storefront with a fully dynamic CMS admin panel. The admin panel allows non-technical "
        "store owners to manage all website content — products, categories, homepage sections, "
        "team members, contact details, SEO settings, orders, and more — without requiring any "
        "code changes. Every piece of content visible on the storefront is configurable from the "
        "admin panel in real time.",
        "Normal", False
    ),

    # ── 4.12 ─────────────────────────────────────────────────────────────────
    ("4.12 Objectives of the Project", "Normal", True),
    ("· To build a complete full-stack furniture e-commerce web application with a public storefront and an admin CMS panel.", "List Paragraph", False),
    ("· To design a dynamic, configuration-driven admin system capable of managing 23 content entities without hard-coding individual forms for each entity.", "List Paragraph", False),
    ("· To implement a MongoDB-backed data persistence layer with automatic seeding of default content on first startup.", "List Paragraph", False),
    ("· To create a responsive, accessible, and visually consistent public storefront driven entirely by CMS-configurable content.", "List Paragraph", False),
    ("· To apply monorepo architecture using pnpm workspaces and Turborepo for structured multi-package project management.", "List Paragraph", False),
    ("· To demonstrate practical full-stack skills including REST API design, component-based UI development, image upload management, server-side validation, and real-time content visibility controls.", "List Paragraph", False),

    # ── 4.13 ─────────────────────────────────────────────────────────────────
    ("4.13 Methodology / Approach Adopted", "Normal", True),
    ("The project followed a feature-driven iterative development approach organized into five phases:", "Normal", False),
    ("· Phase 1 – Architecture and Setup: Established the monorepo structure using pnpm workspaces and Turborepo. Configured the frontend (Vite + React + TypeScript) and backend (Node.js + Express + MongoDB) as separate workspace applications, with a shared @repo/schemas package for cross-layer Zod validation.", "List Paragraph", False),
    ("· Phase 2 – Backend API and Data Layer: Designed the MongoDB data model, implemented the Express REST API, built the dynamic entityConfigs system, and developed server-side validation, normalization, and activity logging logic.", "List Paragraph", False),
    ("· Phase 3 – Admin CMS Panel: Developed the admin layout, overview dashboard, universal entity management component (AdminDashboard.tsx), and the configuration-driven dynamic form renderer (DynamicForm.tsx).", "List Paragraph", False),
    ("· Phase 4 – Public Storefront: Built all public-facing pages (Home, Collections, About, Contact) with full CMS integration, responsive design, and page visibility control driven by admin toggles.", "List Paragraph", False),
    ("· Phase 5 – Integration and Refinement: Connected all layers, implemented image upload and serving, added in-memory API caching on the frontend, and applied performance optimizations including lazy loading and image format optimization.", "List Paragraph", False),

    # ── 4.14 ─────────────────────────────────────────────────────────────────
    ("4.14 Design, Development, and Implementation", "Normal", True),
    ("4.14.1 Technology Stack", "Normal", True),
    ("The following technologies were used in the development of the Bhavika Furniture project:", "Normal", False),
    ("· React.js (v18) with TypeScript – Frontend UI framework for the storefront and admin panel.", "List Paragraph", False),
    ("· Vite – Frontend build tool and development server running on Port 5173.", "List Paragraph", False),
    ("· Tailwind CSS (v4) – Utility-first CSS framework for responsive styling with custom design tokens.", "List Paragraph", False),
    ("· Ant Design (v5) – UI component library for admin panel forms, tables, and input controls.", "List Paragraph", False),
    ("· React Router DOM (v7) – Client-side routing for public and admin pages.", "List Paragraph", False),
    ("· React Hook Form – Form state management and client-side field validation.", "List Paragraph", False),
    ("· React Quill – Rich text editor for product descriptions and long-form page content.", "List Paragraph", False),
    ("· Node.js + Express.js – Backend REST API server running on Port 3000.", "List Paragraph", False),
    ("· MongoDB (Mongoose) – NoSQL database for all application data persistence.", "List Paragraph", False),
    ("· Multer – Image upload middleware configured with in-memory storage.", "List Paragraph", False),
    ("· Zod (@repo/schemas) – Shared validation schemas referenced by both frontend and backend.", "List Paragraph", False),
    ("· pnpm Workspaces + Turborepo – Monorepo management and parallel build/dev orchestration.", "List Paragraph", False),
    ("· Axios – HTTP client for frontend API communication with a 10-second in-memory cache layer.", "List Paragraph", False),
    ("· Lucide React – Consistent icon library used throughout the application.", "List Paragraph", False),

    ("4.14.2 Project Architecture", "Normal", True),
    (
        "The project uses a monorepo structure managed with pnpm workspaces and Turborepo. "
        "The apps/frontend workspace contains the React + TypeScript SPA. The apps/backend "
        "workspace contains the Node.js + Express REST API. The packages/schemas workspace "
        "contains shared Zod validation schemas referenced by both applications using workspace "
        "linking (workspace:*). The turbo.json pipeline enables both applications to start in "
        "parallel with a single turbo run dev command, and the build pipeline respects "
        "inter-package dependencies to ensure correct build order.",
        "Normal", False
    ),

    ("4.14.3 Backend Implementation", "Normal", True),
    (
        "Database Models: Two Mongoose models are defined. The DbState model is a "
        "single-document store that holds all CMS content as a structured JSON object covering "
        "products, categories, orders, customers, page settings, and more. The Image model "
        "stores uploaded images as binary Buffer data with filename, contentType, and automatic "
        "timestamps.",
        "Normal", False
    ),
    (
        "Automatic Data Seeding: On first startup, the initDb() function seeds the database "
        "with default content including products, categories, homepage settings, testimonials, "
        "team members, and contact details. A secondary function ensureCommerceData() adds "
        "commerce entities (orders, customers, coupons, shipping settings, activity log) if "
        "absent, ensuring the application is fully functional immediately after deployment "
        "without any manual data entry.",
        "Normal", False
    ),
    (
        "Mutation Queue Pattern: To prevent race conditions on the shared database document "
        "during concurrent write operations, all database mutations are serialized through a "
        "JavaScript Promise-based mutation queue. Each write operation is chained onto the "
        "previous one, guaranteeing sequential execution without requiring database-level "
        "locking.",
        "Normal", False
    ),
    (
        "Dynamic Entity Configuration (entityConfigs): A central configuration object defines "
        "field schemas for all 23 managed entities. Each entity's configuration specifies field "
        "names, types (text, number, select, rich-text, image, image-gallery, key-value, tags, "
        "datetime, textarea, category-select), labels, and required status. This single object "
        "drives backend validation logic, frontend form generation, and admin table column "
        "rendering simultaneously — eliminating code duplication across the entire stack.",
        "Normal", False
    ),
    (
        "Server-Side Validation: The validatePayload() function enforces required field checks, "
        "number and date format validation, select option constraint enforcement, unique field "
        "constraints (SKUs, slugs, coupon codes, email addresses), and commerce-specific "
        "business rules such as sale price must be lower than regular price, coupon percentage "
        "cannot exceed 100%, and SEO meta title length is limited to 70 characters.",
        "Normal", False
    ),
    (
        "Activity Logging: Every create, update, and delete operation automatically appends an "
        "entry to the activity_log entity recording the action name, entity type, item label, "
        "user, and ISO timestamp. The log is capped at 250 entries to prevent unbounded growth.",
        "Normal", False
    ),

    ("Admin REST API Endpoints (/api/admin):", "Normal", True),
    ("· POST /api/admin/upload – Upload an image; stored as binary Buffer in MongoDB.", "List Paragraph", False),
    ("· GET /api/admin/images/:id – Serve a stored image by MongoDB ObjectId with correct Content-Type header.", "List Paragraph", False),
    ("· GET /api/admin/config/:entity – Fetch field configuration for dynamic UI rendering.", "List Paragraph", False),
    ("· GET /api/admin/data/:entity – Fetch entity data with optional search, status filter, and pagination.", "List Paragraph", False),
    ("· POST /api/admin/data/:entity – Create a new entity record with full server-side validation.", "List Paragraph", False),
    ("· PUT /api/admin/data/:entity/:id – Update an existing entity record.", "List Paragraph", False),
    ("· DELETE /api/admin/data/:entity/:id – Delete an entity record.", "List Paragraph", False),
    ("· POST /api/admin/data/:entity/bulk-status – Bulk activate or hide multiple selected records.", "List Paragraph", False),
    ("· GET /api/admin/dashboard – Aggregated store analytics data.", "List Paragraph", False),
    ("· GET /api/admin/inventory/low-stock – Products with stock at or below 5 units.", "List Paragraph", False),

    ("Storefront REST API Endpoints (/api/store):", "Normal", True),
    ("· GET /api/store/catalog – Fetch all active products and categories for public display.", "List Paragraph", False),
    ("· GET /api/store/content/:entity – Fetch public CMS content for any named entity.", "List Paragraph", False),
    ("· POST /api/store/coupons/validate – Validate a coupon code against all business rules.", "List Paragraph", False),
    ("· POST /api/store/orders – Place a customer order.", "List Paragraph", False),

    ("4.14.4 Frontend Implementation", "Normal", True),
    (
        "The frontend is a React 18 + TypeScript single-page application built with Vite. All "
        "page components are lazy-loaded using React.lazy() and Suspense for optimal initial "
        "bundle size and page load performance. A cohesive design system is implemented using "
        "CSS custom properties: --ink (#171329) for primary text and dark backgrounds, "
        "--violet (#6c4cff) for the primary brand accent and CTAs, --coral (#ff5c4d) for "
        "secondary accents, --citrus (#dfff5b) for highlight elements, and --lavender (#e7ddff) "
        "for light backgrounds and hover states. Typography uses DM Sans for body text and "
        "Manrope for display headings.",
        "Normal", False
    ),
    ("Public Storefront Pages:", "Normal", True),
    (
        "· Home Page – Fetches and renders a CMS-configurable hero section (title, subtitle, "
        "background image, CTA button), a bento-style category grid displaying up to four active "
        "categories, a features/highlights section, and a testimonials panel. All content is "
        "dynamically driven from the admin CMS.",
        "List Paragraph", False
    ),
    (
        "· Category Page – Displays either a directory of all active categories or a filtered "
        "product grid based on the ?type=slug URL parameter. A sticky filter bar with category "
        "thumbnail images allows quick switching between collections. The page fully respects "
        "the admin visibility toggle for the products section.",
        "List Paragraph", False
    ),
    (
        "· About Page – Renders the brand story in rich text, team member profile cards with "
        "photos and bios, and a gallery section — all from CMS data. Active status filtering "
        "ensures only published team members are displayed.",
        "List Paragraph", False
    ),
    (
        "· Contact Page – Displays configurable addresses, phone numbers, email addresses, and "
        "store location cards from the CMS. Includes a contact message form with client-side "
        "submission handling.",
        "List Paragraph", False
    ),
    (
        "· Page Visibility System – The public layout fetches page_visibility configuration "
        "from the CMS on mount. If a page section is toggled off by the admin, the corresponding "
        "route renders a Page Unavailable screen instead of its content, enabling instant "
        "publish/unpublish of entire sections without any code changes.",
        "List Paragraph", False
    ),
    ("Admin Panel:", "Normal", True),
    (
        "· Admin Layout – A fixed 292px sidebar with 22 content entities organized into seven "
        "groups: Business, Website, Catalog, Marketing, About, Contact, and Administration. "
        "Includes a mobile hamburger menu with backdrop overlay, sticky breadcrumb header, and "
        "a contextual editing banner for the active section.",
        "List Paragraph", False
    ),
    (
        "· Admin Overview Dashboard – Displays four key business metrics (paid revenue, open "
        "orders, customer count, low-stock product count), a recent orders feed, a store health "
        "panel showing catalog readiness and order completion percentages as visual progress "
        "bars, a low-stock inventory alert list, and a latest activity log feed.",
        "List Paragraph", False
    ),
    (
        "· Admin Dashboard (Universal Entity Manager) – A single reusable component powering "
        "all 22 entity management sections. Auto-generates data table columns from entity field "
        "configuration, provides search and category filter controls, supports bulk "
        "activate/hide operations, and includes a live page visibility toggle switch per section.",
        "List Paragraph", False
    ),
    (
        "· Dynamic Form Renderer – A fully dynamic form component that renders appropriate "
        "input controls for 13 field types based on configuration fetched from the backend API. "
        "Supported types include: text, number, select, datetime picker, textarea, rich text "
        "(Quill), single image upload, multi-image gallery, key-value specification pairs, "
        "tags (multi-select), and category selector. Integrates with the media library for "
        "image reuse across entities.",
        "List Paragraph", False
    ),

    # ── 4.15 ─────────────────────────────────────────────────────────────────
    ("4.15 Entities Managed by the CMS", "Normal", True),
    ("The admin panel provides full CRUD management for the following 23 entities across seven functional groups:", "Normal", False),
    ("· Business: Orders (full lifecycle tracking from New to Delivered), Customers (profile and lifetime value management).", "List Paragraph", False),
    ("· Website: Home Page settings (hero title, subtitle, image, CTA), Home Features (highlight cards with ordering), Home Testimonials (customer reviews with star ratings).", "List Paragraph", False),
    ("· Catalog: Categories (slug auto-generation, images, starting price, publish scheduling), Products (SKU, rich description, pricing, sale price, stock, image gallery, key-value specifications, tags, publish scheduling).", "List Paragraph", False),
    ("· Marketing: Coupons (percentage/fixed/free-shipping types with usage limits and expiry), Media Library (reusable brand imagery with folder organization), SEO Settings (meta title, description, social image per page).", "List Paragraph", False),
    ("· About: About Page (brand story, mission in rich text, header image), Team Members (photos, roles, bios), Gallery (workshop and showroom imagery).", "List Paragraph", False),
    ("· Contact: Addresses, Phone Numbers, Email Addresses, Store Locations (with photos and full address).", "List Paragraph", False),
    ("· Administration: Shipping Settings (zones, standard rates, free-shipping threshold, delivery estimates), Admin Users (roles and access status), Activity Log (read-only chronological audit trail), Page Visibility (master on/off toggles for all sections).", "List Paragraph", False),

    # ── 4.16 ─────────────────────────────────────────────────────────────────
    ("4.16 Results and Discussion", "Normal", True),
    ("The Bhavika Furniture project was successfully completed as a fully functional full-stack web application. Key outcomes achieved include:", "Normal", False),
    ("· A complete public storefront with four fully CMS-driven pages, fully responsive across mobile and desktop screen sizes.", "List Paragraph", False),
    ("· A dynamic admin CMS panel managing 23 content entities through a single reusable dashboard and form component, eliminating the need to write individual forms or table views for each entity.", "List Paragraph", False),
    ("· A working image upload system that stores images as binary data in MongoDB and serves them through a dedicated API endpoint, with media library reuse supported across all entities.", "List Paragraph", False),
    ("· A real-time page visibility control system allowing entire storefront sections to be published or hidden instantly via admin toggle switches without any code changes or server restarts.", "List Paragraph", False),
    ("· An automatic activity log recording every content change with actor, action, item, and timestamp — providing a complete and chronological audit trail of all workspace activity.", "List Paragraph", False),
    ("· A functional coupon and discount system with configurable discount types (percentage, fixed amount, free shipping), usage limits, minimum order requirements, and expiry date support.", "List Paragraph", False),
    ("· A complete order management workflow supporting the full order lifecycle from New through Processing, Preparing, Shipped, Delivered, and Returned states with fulfilment status tracking.", "List Paragraph", False),
    ("· Frontend performance optimizations including lazy-loaded routes via React.lazy(), a 10-second in-memory API response cache, WebP image format optimization for Unsplash-sourced imagery, and Vite manual chunk splitting for vendor and UI libraries.", "List Paragraph", False),
    (
        "The project demonstrated that a configuration-driven architecture significantly reduces "
        "development effort when building CMS-style applications. The single entityConfigs "
        "object serving as the source of truth for backend validation, frontend form rendering, "
        "and admin table column generation proved to be a highly effective and extensible design "
        "pattern that is directly applicable to the type of enterprise software development "
        "practised at Airwix Technologies.",
        "Normal", False
    ),

    # ── 4.17 ─────────────────────────────────────────────────────────────────
    ("4.17 Challenges Faced and Solutions", "Normal", True),
    (
        "· Single Document Write Concurrency: Storing all application data in a single MongoDB "
        "document introduced the risk of data corruption during concurrent write operations. "
        "A JavaScript Promise-based mutation queue was implemented to serialize all writes, "
        "ensuring each operation completes before the next begins without requiring "
        "database-level locking mechanisms.",
        "List Paragraph", False
    ),
    (
        "· Configuration-Driven Form for 13 Field Types: Building one form component that "
        "dynamically renders appropriate inputs for 13 different field types across 23 entities "
        "required careful abstraction. The solution was a Controller-based render switch in "
        "DynamicForm.tsx where each field type maps to its corresponding Ant Design or custom "
        "input component, keeping the implementation clean and easily extensible.",
        "List Paragraph", False
    ),
    (
        "· Active Status Boolean Inconsistency: The active field was stored as the string "
        "'true'/'false' due to HTML select controls sending string values. All filtering logic "
        "was made consistent by checking for both native boolean === true and string === 'true' "
        "values throughout all frontend and backend code.",
        "List Paragraph", False
    ),
    (
        "· Responsive Admin Sidebar on Mobile: A translate-based CSS toggle pattern was "
        "implemented for the admin sidebar — always present in the DOM but translated off-screen "
        "on mobile with a smooth transition animation. A backdrop overlay prevents interaction "
        "with the main content while the sidebar is open, replicating the behaviour of native "
        "mobile navigation drawers.",
        "List Paragraph", False
    ),
    (
        "· Image Storage and Serving: Multer was configured with memory storage to keep uploaded "
        "files as in-memory Buffers rather than writing to disk. The Buffer is stored in MongoDB "
        "via the Image model, and served back through a dedicated /api/admin/images/:id endpoint "
        "that reads the Buffer and sets the correct Content-Type header based on the stored "
        "MIME type, supporting JPEG, PNG, and WebP formats.",
        "List Paragraph", False
    ),

    # ── 4.18 ─────────────────────────────────────────────────────────────────
    ("4.18 Comparison: ERP & HRMS Project vs. Bhavika Furniture Project", "Normal", True),
    (
        "Both projects undertaken during the internship period complemented each other "
        "significantly. The professional coding standards, REST API design patterns, "
        "component architecture practices, and performance optimization techniques learned "
        "through the ERP and HRMS development work at Airwix Technologies were directly applied "
        "and reinforced through the independent development of the Bhavika Furniture platform.",
        "Normal", False
    ),
    (
        "· ERP & HRMS System (Airwix Technologies): Enterprise team project with a full "
        "development team and senior mentors. Stack: React.js, Node.js, MongoDB, MySQL, JWT "
        "authentication with four-role RBAC. Methodology: Agile-Scrum with two-week sprints, "
        "daily stand-ups, and code reviews. Key learning areas: enterprise SDLC, team "
        "collaboration, production debugging, and multi-module integration.",
        "List Paragraph", False
    ),
    (
        "· Bhavika Furniture E-Commerce CMS (Personal Project): Independent full-stack product built "
        "solo. Stack: React + TypeScript, Vite, Tailwind CSS, Ant Design, Node.js, MongoDB, "
        "Turborepo monorepo. Key learning areas: TypeScript proficiency, monorepo architecture, "
        "configuration-driven UI design, dynamic form rendering, image binary storage, and "
        "in-memory caching strategies.",
        "List Paragraph", False
    ),
    (
        "The key distinction between the two projects is scope and context. The ERP and HRMS "
        "project provided exposure to enterprise-grade professional development practices "
        "including code reviews, sprint planning, and structured team collaboration. The Bhavika Furniture "
        "project reinforced the ability to independently architect, implement, and deliver a "
        "complete production-ready product — a skill set equally essential for a successful "
        "career in full-stack software development.",
        "Normal", False
    ),
]

# ── Insert blocks in REVERSE order so addnext produces correct sequence ───────
for text, style, is_bold in reversed(blocks):
    new_p_elem = OxmlElement('w:p')
    anchor._element.addnext(new_p_elem)
    p = Paragraph(new_p_elem, doc)
    try:
        p.style = doc.styles[style]
    except KeyError:
        p.style = doc.styles['Normal']
    run = p.add_run(text)
    if is_bold:
        run.bold = True

# ── Save updated document ─────────────────────────────────────────────────────
output_path = r'C:\Users\miraj\OneDrive\Desktop\folder\MirajPipaliya_Internship_Report_University_UPDATED.docx'
doc.save(output_path)
print(f"SUCCESS – saved to: {output_path}")
